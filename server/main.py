import os
import json
import re
import unicodedata
import base64
import mammoth
from bs4 import BeautifulSoup
import io

from fastapi import FastAPI, Query, File, UploadFile, Depends, HTTPException
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
from sqlalchemy.orm import Session
from PyPDF2 import PdfReader
import docx
import openpyxl
from fastapi.middleware.cors import CORSMiddleware

from database import SessionLocal, Documento, Usuario, get_db
from groq import Groq

app = FastAPI(title="Sardiñas na Nube Enterprise AI")

MODELO_RAPIDO = "llama-3.1-8b-instant"
MODELO_PRO = "llama-3.3-70b-versatile" 
MODELO_VISION = "llama-3.2-11b-vision-preview"

client_ai = Groq(api_key="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")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
app.mount("/uploads", StaticFiles(directory=UPLOAD_DIR), name="uploads")

def normalizar_texto(texto: str) -> str:
    if not texto: return ""
    texto = unicodedata.normalize("NFD", texto)
    return texto.encode("ascii", "ignore").decode("utf-8").lower()

def categorizar_con_ia(texto: str) -> str:
    if not texto or len(texto) < 10: return "General"
    try:
        prompt = f"Clasifica el siguiente texto OBLIGATORIAMENTE en UNA SOLA PALABRA de esta lista exacta: Finanzas, Legal, RRHH, Proyectos, Software, General. Si ves código de programación (html, js, python, etc.), responde 'Software'. NO des explicaciones. NO uses signos de puntuación. Solo la palabra.\n\nTexto:\n{texto[:1500]}"
        
        res = client_ai.chat.completions.create(
            messages=[{"role": "user", "content": prompt}], 
            model=MODELO_RAPIDO,
            temperature=0.0
        )
        
        cat_raw = res.choices[0].message.content.strip().upper()
        
        if "FINANZAS" in cat_raw: return "Finanzas"
        if "LEGAL" in cat_raw: return "Legal"
        if "RRHH" in cat_raw: return "RRHH"
        if "PROYECTOS" in cat_raw: return "Proyectos"
        if "SOFTWARE" in cat_raw: return "Software"
        
        return "General"
    except: 
        return "General"
    
def reconstruir_docx(html_content, ruta_archivo):
    """Convierte el HTML editado de vuelta a un archivo Word físico con imágenes."""
    try:
        nuevo_doc = docx.Document()
        soup = BeautifulSoup(html_content, 'html.parser')
        
        for element in soup.find_all(['p', 'h1', 'h2', 'h3', 'img', 'li']):
            try:
                if element.name in ['h1', 'h2', 'h3']:
                    nuevo_doc.add_heading(element.get_text(), level=int(element.name[1]))
                    
                elif element.name == 'p':
                    texto = element.get_text().strip()
                    if texto:
                        nuevo_doc.add_paragraph(texto)
                        
                elif element.name == 'li':
                    texto = element.get_text().strip()
                    if texto:
                        nuevo_doc.add_paragraph(texto, style='List Bullet')
                        
                elif element.name == 'img':
                    src = element.get('src', '')
                    if src.startswith('data:image'):
                        head, encoded = src.split(",", 1)
                        img_data = base64.b64decode(encoded)
                        img_stream = io.BytesIO(img_data)
                        nuevo_doc.add_picture(img_stream)
            except Exception as e:
                print(f"Error procesando un elemento para el DOCX: {e}")
                    
        nuevo_doc.save(ruta_archivo)
        print("¡Éxito! Archivo DOCX físico sobrescrito.")
    except Exception as e:
        print(f"Error crítico al reescribir DOCX: {e}")

@app.post("/importar/")
async def importar_documento(file: UploadFile = File(...), db: Session = Depends(get_db)):
    usuario = db.query(Usuario).first()
    if not usuario:
        usuario = Usuario(username="admin", email="admin@demo.com", password_hash="1234")
        db.add(usuario); db.commit(); db.refresh(usuario)

    ruta_guardado = os.path.join(UPLOAD_DIR, file.filename)
    contenido_binario = await file.read()
    with open(ruta_guardado, "wb") as buffer:
        buffer.write(contenido_binario)

    texto_extraido = ""
    excel_grid = []
    docx_html_preview = ""
    extension = file.filename.split(".")[-1].lower()

    try:
        if extension == "pdf":
            reader = PdfReader(ruta_guardado)
            texto_extraido = "\n".join([page.extract_text() or "" for page in reader.pages])
        
        elif extension in ["doc", "docx"]: 
            with open(ruta_guardado, "rb") as docx_file:
                result = mammoth.convert_to_html(docx_file)
                html_contenido = result.value 
                
                docx_file.seek(0)
                
                texto_extraido = mammoth.extract_raw_text(docx_file).value
                docx_html_preview = html_contenido
                
        elif extension in ["xls", "xlsx"]:
            wb = openpyxl.load_workbook(ruta_guardado, data_only=True)
            for sheet in wb.worksheets:
                filas = []
                for row in sheet.iter_rows(values_only=True):
                    fila_str = [str(c) if c is not None else "" for c in row]
                    if any(fila_str):
                        filas.append(fila_str)
                        texto_extraido += " | ".join(fila_str) + "\n"
                excel_grid.append({"hoja": sheet.title, "filas": filas})
        elif extension in ["txt", "md", "csv", "json", "py", "js", "html", "css"]:
            try:
                texto_extraido = contenido_binario.decode('utf-8')
                
                if extension == "csv":
                    texto_extraido = texto_extraido.replace(",", " | ")
                
            except UnicodeDecodeError:
                texto_extraido = contenido_binario.decode('latin-1', errors='ignore')
        elif extension in ["jpg", "jpeg", "png"]:
            b64_img = base64.b64encode(contenido_binario).decode('utf-8')
            res_v = client_ai.chat.completions.create(
                messages=[{"role": "user", "content": [{"type": "text", "text": "Lee todo el texto de esta imagen."}, {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{b64_img}"}}]}],
                model=MODELO_VISION
            )
            texto_extraido = res_v.choices[0].message.content
    except Exception as e:
        texto_extraido = f"Error: {e}"

    categoria = categorizar_con_ia(texto_extraido)
    metadatos = {"extension": extension, "categoria": categoria, "excel_grid": excel_grid, "html_preview": docx_html_preview}
    
    nuevo_doc = Documento(
        nombre_archivo=file.filename,
        nombre_normalizado=normalizar_texto(file.filename),
        ruta_archivo=ruta_guardado,
        contenido_texto=texto_extraido,
        contenido_normalizado=normalizar_texto(texto_extraido),
        metadatos=json.dumps(metadatos),
        user_id=usuario.id
    )
    db.add(nuevo_doc); db.commit()
    return {"mensaje": "OK"}

@app.put("/documentos/{doc_id}")
async def actualizar_documento(doc_id: int, payload: dict, db: Session = Depends(get_db)):
    doc = db.query(Documento).filter(Documento.id == doc_id).first()
    m = json.loads(doc.metadatos)
    
    if "html_preview" in payload:
        m["html_preview"] = payload["html_preview"]
        if m.get("extension", "").lower() in ["doc", "docx"]:
            reconstruir_docx(payload["html_preview"], doc.ruta_archivo)

    if "excel_grid" in payload:
        m["excel_grid"] = payload["excel_grid"]
        nuevo_texto = ""
        for hoja in payload["excel_grid"]:
            for fila in hoja.get("filas", []):
                nuevo_texto += " | ".join([str(c) for c in fila]) + "\n"
        doc.contenido_texto = nuevo_texto
        doc.contenido_normalizado = normalizar_texto(nuevo_texto)
        m["categoria"] = categorizar_con_ia(nuevo_texto)

    elif "contenido_texto" in payload:
        doc.contenido_texto = payload["contenido_texto"]
        doc.contenido_normalizado = normalizar_texto(payload["contenido_texto"])
        m["categoria"] = categorizar_con_ia(payload["contenido_texto"])

    if "nombre_archivo" in payload:
        doc.nombre_archivo = payload["nombre_archivo"]
        doc.nombre_normalizado = normalizar_texto(payload["nombre_archivo"])
        
    doc.metadatos = json.dumps(m)
    db.commit()
    return {"status": "success", "nueva_categoria": m["categoria"]}

@app.get("/documentos/")
def buscar_y_navegar(query: str = None, categoria: str = None, skip: int = 0, limit: int = 10, db: Session = Depends(get_db)):
    consulta = db.query(Documento)
    
    if query:
        q = normalizar_texto(query)
        consulta = consulta.filter((Documento.nombre_normalizado.like(f"%{q}%")) | (Documento.contenido_normalizado.like(f"%{q}%")))
    
    if categoria and categoria != "Todas":
        consulta = consulta.filter(Documento.metadatos.like(f'%"{categoria}"%'))
    
    res = consulta.offset(skip).limit(limit).all()
    final = []
    
    for d in res:
        m = json.loads(d.metadatos)
        frags = []
        matches = [] 

        if query and d.contenido_texto:
            clean_o = d.contenido_texto.replace("\n", " ")
            clean_n = normalizar_texto(clean_o)
            matches = list(re.finditer(re.escape(normalizar_texto(query)), clean_n))
            
            last_end = -1
            for match in matches:
                start = max(0, match.start() - 50)
                end = min(len(clean_o), match.end() + 50)
                
                if start <= last_end:
                    if frags:
                        frags[-1] = clean_o[max(0, clean_n.find(frags[-1])):end]
                        last_end = end
                else:
                    frags.append(clean_o[start:end])
                    last_end = end
        fecha_str = d.fecha_subida.strftime("%d/%m/%Y") if d.fecha_subida else "Desconocida"
                
        final.append({
            "id": d.id, 
            "nombre_archivo": d.nombre_archivo, 
            "metadatos": m, 
            "fragmentos": frags, 
            "coincidencias": len(matches), 
            "texto_completo": d.contenido_texto,
            "autor": d.propietario.username if d.propietario else "Sistema",
            "fecha": fecha_str
        })
        
    return {"total_encontrados": consulta.count(), "resultados": final}

@app.get("/documentos/{doc_id}/resumir")
async def resumir_documento(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Documento).filter(Documento.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    
    metadatos = json.loads(doc.metadatos)
    ext = metadatos.get("extension", "").lower()

    if ext in ["py", "js", "html", "css", "json"]:
        contexto_ia = (
            f"Analiza este archivo de CÓDIGO o CONFIGURACIÓN (.{ext}). "
            "Explica: 1. Su propósito principal. 2. Las tecnologías o estructuras que usa. "
            "3. Una breve conclusión técnica. Usa emojis técnicos como 💻, ⚙️ o 🚀."
        )
    elif ext in ["csv", "xlsx", "xls"]:
        contexto_ia = (
            "Analiza estos DATOS o TABLA. Resume de qué trata el conjunto de datos, "
            "qué columnas o información relevante contiene y una conclusión de los datos. "
            "Usa emojis como 📊, 📈 o 📁."
        )
    else:
        contexto_ia = (
            "Resume este DOCUMENTO CORPORATIVO en 3 puntos clave de forma ejecutiva. "
            "Usa emojis relevantes según el contenido (⚖️, 💰, 📅, etc.)."
        )

    res = client_ai.chat.completions.create(
        messages=[
            {
                "role": "system", 
                "content": "Eres un experto en análisis de documentos y código técnico. Tu objetivo es ser breve y preciso."
            },
            {
                "role": "user", 
                "content": f"{contexto_ia}\n\nContenido del archivo:\n{doc.contenido_texto[:3500]}"
            }
        ], 
        model=MODELO_PRO
    )
    
    return {"resumen": res.choices[0].message.content}

@app.get("/documentos/{doc_id}/descargar")
async def descargar_documento(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Documento).filter(Documento.id == doc_id).first()
    return FileResponse(path=doc.ruta_archivo, filename=doc.nombre_archivo)

@app.delete("/documentos/{doc_id}")
async def eliminar_documento(doc_id: int, db: Session = Depends(get_db)):
    doc = db.query(Documento).filter(Documento.id == doc_id).first()
    if os.path.exists(doc.ruta_archivo): os.remove(doc.ruta_archivo)
    db.delete(doc); db.commit()
    return {"mensaje": "OK"}